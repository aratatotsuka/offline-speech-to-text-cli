from __future__ import annotations

from pathlib import Path

from docx import Document

from app.errors import DocxConversionError


def txt_to_docx(*, txt_path: Path, docx_path: Path, title: str) -> None:
    if not txt_path.exists():
        raise DocxConversionError(f"txt not found: {txt_path}")

    try:
        text = txt_path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        raise DocxConversionError(f"failed to read txt: {txt_path} ({exc})")

    try:
        doc = Document()
        doc.add_heading(title, level=1)
        for line in text.splitlines():
            doc.add_paragraph(line)
        doc.save(docx_path)
    except Exception as exc:
        raise DocxConversionError(f"failed to write docx: {docx_path} ({exc})")

